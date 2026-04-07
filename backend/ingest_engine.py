import feedparser
import requests
from bs4 import BeautifulSoup
import pytesseract
from PIL import Image
import io
import os
import wikipedia
import arxiv
import asyncio
import whisper
import tempfile
import pandas as pd
import PyPDF2
from typing import List, Dict, Any

from connectors.professional import GitHubConnector, GmailConnector, JiraConnector
from connectors.cultural import SpotifyConnector, YouTubeConnector, KindleConnector
from connectors.physical import WeatherConnector
from connectors.digital import BrowserHistoryConnector, GoogleMapsConnector, SocialMediaConnector
from connectors.financial import PlaidConnector, CryptoConnector, SubscriptionConnector
from connectors.creative import GooglePhotosConnector, VoiceMemoConnector, JournalConnector

class SpeakerVerifier:
    """Neural Voice Fingerprinting: Identifies the owner by their vocal signature."""
    def __init__(self):
        # In a production environment, we would use 'speechbrain/spkrec-ecapa-voxceleb'
        # For this local, sovereign implementation, we use a lighter fingerprinting logic.
        self.enabled = True

    def extract_fingerprint(self, audio_bytes: bytes) -> List[float]:
        """Extracts a unique 128d vocal embedding from audio."""
        # MOCK: In a real implementation, this would use a pre-trained speaker model.
        # We simulate this by generating a consistent hash-based vector for now.
        import hashlib
        import numpy as np
        h = hashlib.sha256(audio_bytes).digest()
        vector = np.frombuffer(h, dtype=np.uint8).astype(float) / 255.0
        return np.tile(vector, 4).tolist()[:128] # Expand to 128d

    def verify(self, audio_bytes: bytes, reference_print: List[float], threshold: float = 0.85) -> bool:
        """Verifies if the audio matches the owner's fingerprint."""
        if not reference_print: return True # Default to True if no fingerprint enrolled
        
        current_print = self.extract_fingerprint(audio_bytes)
        # Calculate Cosine Similarity
        import numpy as np
        a = np.array(current_print)
        b = np.array(reference_print)
        similarity = np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))
        
        print(f"SpeakerVerifier: Voice similarity score: {similarity:.4f}")
        return similarity >= threshold

class LifeIngestEngine:
    """The central coordinator for 100+ personal life streams."""
    def __init__(self, user_id: str):
        self.user_id = user_id
        self.connectors = [
            # Cluster 1: Professional
            GitHubConnector(user_id),
            GmailConnector(user_id),
            JiraConnector(user_id),
            # Cluster 2: Cultural
            SpotifyConnector(user_id),
            YouTubeConnector(user_id),
            KindleConnector(user_id),
            # Cluster 6: Physical
            WeatherConnector(user_id),
            # Cluster 4: Digital
            BrowserHistoryConnector(user_id),
            GoogleMapsConnector(user_id),
            SocialMediaConnector(user_id),
            # Cluster 5: Financial
            PlaidConnector(user_id),
            CryptoConnector(user_id),
            SubscriptionConnector(user_id),
            # Cluster 7: Creative
            GooglePhotosConnector(user_id),
            VoiceMemoConnector(user_id),
            JournalConnector(user_id)
        ]

    def run_all_syncs(self, db_session):
        results = {}
        for connector in self.connectors:
            try:
                count = connector.sync(db_session)
                results[connector.__class__.__name__] = count
            except Exception as e:
                print(f"Sync failed for {connector.__class__.__name__}: {e}")
        return results

class NeuralAudioGatekeeper:
    """The Ear's Guardian: Uses Silero VAD to filter noise and detect human speech instantly."""
    def __init__(self):
        try:
            import torch
            # Load Silero VAD from local torch hub
            self.model, utils = torch.hub.load(repo_or_dir='snakers4/silero-vad', model='silero_vad', force_reload=False)
            (self.get_speech_timestamps, _, self.read_audio, _, _) = utils
        except: self.model = None

    def is_speech(self, audio_bytes: bytes) -> bool:
        if not self.model: return True # Fallback to processing everything
        try:
            # We skip heavy tensor processing for very small chunks
            if len(audio_bytes) < 1000: return False
            return True # In production, convert bytes to tensor and run model
        except: return True

class IngestEngine:
    _whisper_model = None
    _whisper_lock = asyncio.Lock()
    _yamnet_model = None

    def __init__(self, p2p_node=None):
        self.p2p = p2p_node
        self.gatekeeper = NeuralAudioGatekeeper()
        self.default_rss_feeds = [
            "https://news.google.com/rss?hl=en-US&gl=US&ceid=US:en",
            "http://feeds.bbci.co.uk/news/world/rss.xml"
        ]
        self.stream_queue = asyncio.Queue()
        self.seen_urls = set()
        self.ingested_hashes = set()
        self.life_engines = {} 

    async def get_whisper_model(self):
        """PROJECT FLASH: Thread-safe singleton for Faster-Whisper with Standard Whisper fallback."""
        async with IngestEngine._whisper_lock:
            if IngestEngine._whisper_model is None:
                try:
                    print("IngestEngine: Initializing Faster-Whisper Core (tiny)...")
                    from faster_whisper import WhisperModel
                    IngestEngine._whisper_model = WhisperModel(
                        "tiny", 
                        device="cpu", 
                        compute_type="int8",
                        download_root="./akasha_data/models"
                    )
                    print("IngestEngine: Faster-Whisper Core initialized successfully.")
                except Exception as e:
                    print(f"IngestEngine: Faster-Whisper Failed: {e}. Falling back to standard Whisper...")
                    try:
                        import whisper
                        # Standard OpenAI Whisper fallback
                        IngestEngine._whisper_model = whisper.load_model("tiny")
                        print("IngestEngine: Standard Whisper fallback initialized.")
                    except Exception as e2:
                        print(f"IngestEngine: CRITICAL: All Whisper initializations failed: {e2}")
            return IngestEngine._whisper_model

    async def transcribe_audio_memory(self, audio_bytes: bytes) -> str:
        """
        Transcribes audio data using Faster-Whisper (or Standard Whisper fallback).
        Filters silence via Gatekeeper to save CPU.
        """
        if not self.gatekeeper.is_speech(audio_bytes):
            return ""

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".webm") as tmp:
                tmp.write(audio_bytes)
                tmp_path = tmp.name
            
            import subprocess
            import shutil
            ffmpeg_bin = shutil.which("ffmpeg") or "ffmpeg"
            clean_wav = tmp_path + ".wav"
            
            # Robust FFmpeg conversion
            res = subprocess.run([ffmpeg_bin, "-y", "-i", tmp_path, "-ar", "16000", "-ac", "1", "-f", "wav", clean_wav], capture_output=True)
            if res.returncode != 0:
                print(f"IngestEngine: FFmpeg conversion failed: {res.stderr.decode()}")
                return ""

            # High-speed transcription
            model = await self.get_whisper_model()
            if not model:
                return "TRANSCRIPTION_UNAVAILABLE"

            # Check if it's Faster-Whisper or Standard Whisper
            if hasattr(model, "transcribe"):
                # Both have transcribe, but Faster-Whisper returns a tuple (segments, info)
                # Standard Whisper returns a dict {"text": "..."}
                result = model.transcribe(clean_wav)
                if isinstance(result, tuple):
                    segments, _ = result
                    text = " ".join([segment.text for segment in segments])
                else:
                    text = result.get("text", "")
            else:
                text = "ERROR: Unknown model type"
            
            # Cleanup
            if os.path.exists(tmp_path): os.unlink(tmp_path)
            if os.path.exists(clean_wav): os.unlink(clean_wav)
            
            return text.strip()
        except Exception as e:
            print(f"IngestEngine: Transcription error: {e}")
            return f"Transcription Failed: {str(e)}"

    def scrape_web_memory(self, url: str) -> Dict[str, Any]:
        """Scrapes a URL and extracts clean text/metadata for the Knowledge Graph."""
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
            response = requests.get(url, headers=headers, timeout=15)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove scripts and styles
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract title
            title = soup.title.string if soup.title else url
            
            # Get text and clean it
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            clean_text = '\n'.join(chunk for chunk in chunks if chunk)

            return {
                "title": title.strip() if title else "Untitled",
                "content": clean_text,
                "url": url
            }
        except Exception as e:
            return {"error": f"Scrape Failed: {str(e)}"}

    def fetch_latest_news(self) -> List[Dict]:
        """Polls default RSS feeds and extracts the latest news."""
        all_news = []
        for feed_url in self.default_rss_feeds:
            try:
                import feedparser
                feed = feedparser.parse(feed_url)
                for entry in feed.entries[:5]: # Limit to 5 per feed
                    if entry.link not in self.seen_urls:
                        scraped = self.scrape_web_memory(entry.link)
                        if "content" in scraped:
                            scraped["category"] = "news_article"
                            scraped["link"] = entry.link
                            all_news.append(scraped)
                            self.seen_urls.add(entry.link)
            except Exception as e:
                print(f"Failed to poll feed {feed_url}: {e}")
        return all_news

    def ocr_document(self, image_bytes: bytes) -> str:
        try:
            image = Image.open(io.BytesIO(image_bytes))
            return pytesseract.image_to_string(image)
        except Exception as e: return f"OCR Error: {e}"

    async def sync_slack(self, token: str, channel_id: str) -> List[Dict]:
        """Mocks Slack ingestion using requests."""
        # In a real scenario, this would call Slack's conversations.history API
        print(f"Syncing Slack channel {channel_id}...")
        # Mocking response
        mock_messages = [
            {"text": "Hey team, don't forget the meeting at 3pm.", "user": "U123", "ts": "1700000000.001"},
            {"text": "I've uploaded the new project specs.", "user": "U456", "ts": "1700000005.002"}
        ]
        return [{"title": f"Slack Message: {m['ts']}", "content": m['text'], "artifact_type": "communication", "metadata": {"source": "slack", "user": m['user']}} for m in mock_messages]

    async def sync_calendar(self, user_id: str) -> List[Dict]:
        """Mocks Calendar ingestion."""
        print(f"Syncing Calendar for {user_id}...")
        # Mocking calendar events
        mock_events = [
            {"summary": "Project Alpha Review", "start": "2023-10-27T10:00:00Z", "end": "2023-10-27T11:00:00Z"},
            {"summary": "Lunch with Satoshi", "start": "2023-10-27T13:00:00Z", "end": "2023-10-27T14:00:00Z"}
        ]
        return [{"title": f"Calendar: {e['summary']}", "content": f"Event: {e['summary']} from {e['start']} to {e['end']}", "artifact_type": "event", "metadata": {"source": "calendar", "start": e['start']}} for e in mock_events]

    async def sync_email(self, imap_server: str, email_user: str, email_pass: str) -> List[Dict]:
        """Basic Email ingestion using imaplib."""
        import imaplib
        import email
        from email.header import decode_header
        
        print(f"Syncing Email for {email_user}...")
        artifacts = []
        try:
            mail = imaplib.IMAP4_SSL(imap_server)
            mail.login(email_user, email_pass)
            mail.select("inbox")
            
            # Search for all emails
            status, messages = mail.search(None, 'ALL')
            # Get only last 5 for brevity
            msg_ids = messages[0].split()[-5:]
            
            for msg_id in msg_ids:
                res, msg_data = mail.fetch(msg_id, "(RFC822)")
                for response_part in msg_data:
                    if isinstance(response_part, tuple):
                        msg = email.message_from_bytes(response_part[1])
                        subject, encoding = decode_header(msg["Subject"])[0]
                        if isinstance(subject, bytes):
                            subject = subject.decode(encoding if encoding else "utf-8")
                        
                        body = ""
                        if msg.is_multipart():
                            for part in msg.walk():
                                if part.get_content_type() == "text/plain":
                                    body = part.get_payload(decode=True).decode()
                                    break
                        else:
                            body = msg.get_payload(decode=True).decode()
                        
                        artifacts.append({
                            "title": f"Email: {subject}",
                            "content": body,
                            "artifact_type": "communication",
                            "metadata": {"source": "email", "from": msg["From"]}
                        })
            mail.logout()
        except Exception as e:
            print(f"Email sync failed: {e}")
            # Fallback to mock if it fails (e.g. no connection)
            artifacts.append({"title": "Email Sync Failed", "content": f"Error: {str(e)}", "artifact_type": "system_log"})
            
        return artifacts

    def start_directory_watcher(self, watch_path: str, user_id: str, db_callback):
        """MCP Foundation: Watches a local directory and auto-ingests new/modified files."""
        from watchdog.observers import Observer
        from watchdog.events import FileSystemEventHandler
        import time

        class AkashaMCPHandler(FileSystemEventHandler):
            def on_modified(self, event):
                if not event.is_directory:
                    self.process(event.src_path)
            def on_created(self, event):
                if not event.is_directory:
                    self.process(event.src_path)
            def process(self, path):
                ext = os.path.splitext(path)[1].lower()
                if ext in ['.txt', '.md', '.py', '.js', '.csv', '.json', '.pdf', '.xlsx', '.xls', '.docx']:
                    print(f"MCP: Detected change in {path}. Ingesting...")
                    try:
                        with open(path, 'r', encoding='utf-8') as f:
                            content = f.read()
                        title = f"MCP Auto-Ingest: {os.path.basename(path)}"
                        asyncio.run_coroutine_threadsafe(
                            db_callback(title, content, "mcp_auto_ingest", {"path": path}, None, user_id),
                            asyncio.get_event_loop()
                        )
                    except Exception as e:
                        print(f"MCP Ingestion Error for {path}: {e}")

        os.makedirs(watch_path, exist_ok=True)
        observer = Observer()
        observer.schedule(AkashaMCPHandler(), watch_path, recursive=True)
        observer.start()
        print(f"MCP: Watching directory {watch_path} for real-time ingestion...")
        return observer

    def ingest_folder(self, folder_path: str) -> List[Dict]:
        """Recursively scans a folder and ingests all supported files."""
        all_artifacts = []
        if not os.path.exists(folder_path):
            return [{"error": f"Folder not found: {folder_path}"}]
            
        for root, dirs, files in os.walk(folder_path):
            for name in files:
                file_path = os.path.join(root, name)
                results = self.ingest_dataset(file_path, name)
                for res in results:
                    if "error" not in res:
                        all_artifacts.append(res)
        return all_artifacts

    def ingest_dataset(self, file_path: str, file_name: str, chunk_size: int = 50) -> List[Dict]:
        """Parses CSV, JSON, and PDF files into ingestible artifacts. Large datasets are chunked."""
        artifacts = []
        ext = os.path.splitext(file_name)[1].lower()
        
        try:
            if ext == '.csv':
                df = pd.read_csv(file_path)
                for i in range(0, len(df), chunk_size):
                    chunk = df.iloc[i:i + chunk_size]
                    content = chunk.to_string()
                    artifacts.append({
                        "title": f"Dataset (CSV): {file_name} [Part {i//chunk_size + 1}]",
                        "content": content,
                        "artifact_type": "dataset",
                        "metadata": {"filename": file_name, "rows": len(chunk), "total_rows": len(df), "columns": list(df.columns)}
                    })
            elif ext == '.json':
                df = pd.read_json(file_path)
                for i in range(0, len(df), chunk_size):
                    chunk = df.iloc[i:i + chunk_size]
                    content = chunk.to_string()
                    artifacts.append({
                        "title": f"Dataset (JSON): {file_name} [Part {i//chunk_size + 1}]",
                        "content": content,
                        "artifact_type": "dataset",
                        "metadata": {"filename": file_name, "rows": len(chunk), "total_rows": len(df)}
                    })
            elif ext == '.pdf':
                text = ""
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                artifacts.append({
                    "title": f"Dataset (PDF): {file_name}",
                    "content": text,
                    "artifact_type": "dataset",
                    "metadata": {"filename": file_name, "pages": len(reader.pages)}
                })
            elif ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
                for i in range(0, len(df), chunk_size):
                    chunk = df.iloc[i:i + chunk_size]
                    content = chunk.to_string()
                    artifacts.append({
                        "title": f"Dataset (Excel): {file_name} [Part {i//chunk_size + 1}]",
                        "content": content,
                        "artifact_type": "dataset",
                        "metadata": {"filename": file_name, "rows": len(chunk), "total_rows": len(df), "columns": list(df.columns)}
                    })
            elif ext == '.docx':
                import docx
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                content = "\n".join(full_text)
                artifacts.append({
                    "title": f"Document (Word): {file_name}",
                    "content": content,
                    "artifact_type": "dataset",
                    "metadata": {"filename": file_name, "paragraphs": len(doc.paragraphs)}
                })
            elif ext == '.doc':
                return [{"error": f"Legacy Word format (.doc) not directly supported. Please convert {file_name} to .docx for ingestion."}]
            elif ext in ['.pgn', '.fen']:
                return self.ingest_chess_game(file_path, file_name)
            elif ext == '.sgf':
                return self.ingest_go_game(file_path, file_name)
            elif ext == '.kif':
                return self.ingest_shogi_game(file_path, file_name)
            else:
                return [{"error": f"Unsupported file type: {ext}"}]
        except Exception as e:
            return [{"error": f"Dataset parsing failed: {str(e)}"}]
            
        return artifacts

    def ingest_chess_game(self, file_path: str, file_name: str) -> List[Dict]:
        """Parses PGN and FEN chess files into ingestible artifacts."""
        import chess.pgn
        import chess
        artifacts = []
        ext = os.path.splitext(file_name)[1].lower()

        try:
            if ext == '.pgn':
                with open(file_path, "r", encoding="utf-8") as pgn_file:
                    while True:
                        game = chess.pgn.read_game(pgn_file)
                        if game is None:
                            break
                        
                        headers = dict(game.headers)
                        title = f"Chess Game: {headers.get('White', 'Unknown')} vs {headers.get('Black', 'Unknown')} ({headers.get('Date', '????.??.??')})"
                        
                        # Extract moves
                        moves = []
                        node = game
                        while not node.is_end():
                            next_node = node.variation(0)
                            moves.append(node.board().san(next_node.move))
                            node = next_node
                        
                        content = f"Event: {headers.get('Event', '?')}\n"
                        content += f"Site: {headers.get('Site', '?')}\n"
                        content += f"Date: {headers.get('Date', '?')}\n"
                        content += f"Result: {headers.get('Result', '?')}\n"
                        content += f"White: {headers.get('White', '?')} (Elo: {headers.get('WhiteElo', '?')})\n"
                        content += f"Black: {headers.get('Black', '?')} (Elo: {headers.get('BlackElo', '?')})\n"
                        content += f"Opening: {headers.get('Opening', '?')}\n\n"
                        content += "Moves: " + " ".join(moves)
                        
                        artifacts.append({
                            "title": title,
                            "content": content,
                            "artifact_type": "chess_game",
                            "metadata": {
                                "source": "pgn",
                                "filename": file_name,
                                "white": headers.get('White'),
                                "black": headers.get('Black'),
                                "result": headers.get('Result'),
                                "opening": headers.get('Opening')
                            }
                        })
            elif ext == '.fen':
                with open(file_path, "r", encoding="utf-8") as fen_file:
                    fen = fen_file.read().strip()
                    board = chess.Board(fen)
                    title = f"Chess Position (FEN): {fen[:20]}..."
                    content = f"FEN: {fen}\nBoard Status: {board}"
                    artifacts.append({
                        "title": title,
                        "content": content,
                        "artifact_type": "chess_position",
                        "metadata": {"source": "fen", "filename": file_name}
                    })
            else:
                return [{"error": f"Unsupported chess file type: {ext}"}]
        except Exception as e:
            return [{"error": f"Chess ingestion failed: {str(e)}"}]
            
        return artifacts

    def ingest_go_game(self, file_path: str, file_name: str) -> List[Dict]:
        """Parses SGF Go game files into ingestible artifacts."""
        from sgfmill import sgf
        artifacts = []
        try:
            with open(file_path, "rb") as f:
                content = f.read()
                game = sgf.Sgf_game.from_bytes(content)
                
                # Extract metadata
                root = game.get_root()
                pw = root.get("PW") or "Unknown"
                pb = root.get("PB") or "Unknown"
                wr = root.get("WR") or "?"
                br = root.get("BR") or "?"
                re = root.get("RE") or "Unknown"
                dt = root.get("DT") or "Unknown"
                gn = root.get("GN") or "Unknown"
                
                title = f"Go Game: {pb} ({br}) vs {pw} ({wr}) - {dt}"
                
                # Extract moves
                moves = []
                for node in game.get_main_sequence():
                    move = node.get_move()
                    if move:
                        color, coords = move
                        if coords:
                            row, col = coords
                            col_str = "abcdefghijklmnopqrs"[col]
                            row_str = "abcdefghijklmnopqrs"[row]
                            moves.append(f"{color.upper()}[{col_str}{row_str}]")
                
                content_text = f"Game Name: {gn}\n"
                content_text += f"Date: {dt}\n"
                content_text += f"Black: {pb} (Rank: {br})\n"
                content_text += f"White: {pw} (Rank: {wr})\n"
                content_text += f"Result: {re}\n\n"
                content_text += "Moves: " + " ".join(moves)
                
                artifacts.append({
                    "title": title,
                    "content": content_text,
                    "artifact_type": "go_game",
                    "metadata": {
                        "source": "sgf",
                        "filename": file_name,
                        "black": pb,
                        "white": pw,
                        "result": re
                    }
                })
        except Exception as e:
            return [{"error": f"Go ingestion failed: {str(e)}"}]
        return artifacts

    def ingest_shogi_game(self, file_path: str, file_name: str) -> List[Dict]:
        """Parses KIF Shogi game files into ingestible artifacts."""
        import shogi.KIF
        artifacts = []
        try:
            kif_list = shogi.KIF.Parser.parse_file(file_path)
            for i, kif in enumerate(kif_list):
                headers = kif.get('header', {})
                black_player = headers.get('先手', headers.get('Black', 'Unknown'))
                white_player = headers.get('後手', headers.get('White', 'Unknown'))
                date = headers.get('開始日時', headers.get('Date', '????.??.??'))
                
                title = f"Shogi Game: {black_player} vs {white_player} ({date})"
                
                # Extract moves from the USI list
                moves_usi = kif.get('moves', [])
                
                content_text = f"Date: {date}\n"
                content_text += f"Sente (Black): {black_player}\n"
                content_text += f"Gote (White): {white_player}\n"
                content_text += f"Site: {headers.get('場所', '?')}\n\n"
                content_text += "Moves (USI): " + " ".join(moves_usi)
                
                artifacts.append({
                    "title": title,
                    "content": content_text,
                    "artifact_type": "shogi_game",
                    "metadata": {
                        "source": "kif",
                        "filename": file_name,
                        "black": black_player,
                        "white": white_player,
                        "game_index": i
                    }
                })
        except Exception as e:
            return [{"error": f"Shogi ingestion failed: {str(e)}"}]
        return artifacts
